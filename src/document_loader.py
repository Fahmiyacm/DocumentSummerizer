from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document


def load_document(file_path):
    """
    Load PDF, DOCX, TXT files and convert them
    into LangChain Document format.
    """

    try:

        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(
                f"File not found: {file_path}"
            )


        extension = file_path.suffix.lower()


        documents = []


        # -----------------------------
        # PDF Loader
        # -----------------------------

        if extension == ".pdf":

            reader = PdfReader(
                str(file_path)
            )

            for page_number, page in enumerate(
                reader.pages
            ):

                text = page.extract_text()

                if text and text.strip():

                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source": file_path.name,
                                "page": page_number + 1,
                                "file_type": "PDF"
                            }
                        )
                    )


        # -----------------------------
        # DOCX Loader
        # -----------------------------

        elif extension == ".docx":

            docx_file = DocxDocument(
                str(file_path)
            )

            text = "\n".join(
                paragraph.text
                for paragraph in docx_file.paragraphs
                if paragraph.text.strip()
            )


            if text:

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": file_path.name,
                            "page": 1,
                            "file_type": "DOCX"
                        }
                    )
                )


        # -----------------------------
        # TXT Loader
        # -----------------------------

        elif extension == ".txt":

            with open(
                file_path,
                "r",
                encoding="utf-8"
            ) as file:

                text = file.read()


            if text.strip():

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": file_path.name,
                            "page": 1,
                            "file_type": "TXT"
                        }
                    )
                )


        else:

            raise ValueError(
                "Unsupported file format. "
                "Only PDF, DOCX, TXT are supported."
            )


        return documents


    except Exception as error:

        raise RuntimeError(
            f"Document loading failed: {error}"
        )